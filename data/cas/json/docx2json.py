import os
import json
import re
from docx import Document
import argparse

def clean_text(text):
    """
    Cleans text by:
    - Replacing newlines with space
    - Stripping whitespace
    - Removing trailing footnote letters from numeric values (e.g., 6.202a → 6.202)
    """
    text = text.replace('\n', ' ').strip()

    # Remove footnote letter if it follows a number (e.g. 6.202a → 6.202)
    text = re.sub(r'(?<=\d)[a-zA-Z]$', '', text)

    return text

def extract_single_table_from_docx(docx_path):
    """
    Extracts the only table from a DOCX file (assumes there is exactly one table).
    Returns a dictionary with headers and data rows.
    """
    document = Document(docx_path)

    if not document.tables:
        raise ValueError("No tables found in the document.")

    table = document.tables[0]
    rows = table.rows
    if not rows:
        raise ValueError("Table is empty.")

    headers = [clean_text(cell.text) for cell in rows[0].cells]
    data_rows = []

    for row in rows[1:]:
        cells = [clean_text(cell.text) for cell in row.cells]
        row_dict = {headers[i]: cells[i] if i < len(cells) else "" for i in range(len(headers))}
        data_rows.append(row_dict)

    return {
        "headers": headers,
        "data": data_rows
    }

def process_directory(input_dir, output_dir):
    """
    Processes all .docx files in a directory and writes corresponding .json files (one per input).
    """
    os.makedirs(output_dir, exist_ok=True)

    for filename in os.listdir(input_dir):
        if filename.lower().endswith(".docx"):
            input_path = os.path.join(input_dir, filename)
            base_name = os.path.splitext(filename)[0]
            output_path = os.path.join(output_dir, base_name + ".json")

            try:
                table = extract_single_table_from_docx(input_path)
                with open(output_path, "w", encoding="utf-8") as f:
                    json.dump(table, f, ensure_ascii=False, indent=2)
                print(f"✅ Converted: {filename} → {output_path}")
            except Exception as e:
                print(f"❌ Failed to process {filename}: {e}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Convert single-table DOCX files in a directory to cleaned JSON.")
    parser.add_argument("--input-dir", "-i", required=True, help="Directory containing .docx files")
    parser.add_argument("--output-dir", "-o", default=".", help="Directory to save .json files")
    args = parser.parse_args()

    process_directory(args.input_dir, args.output_dir)
